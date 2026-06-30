from __future__ import annotations

import os
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.models.documents import RawDocument
from app.utils.logger import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({"pdf", "docx", "txt", "xlsx"})


class UnsupportedFileTypeError(Exception):
    pass


class DocumentLoader:
    def load(
        self,
        file_path: str | Path,
        original_filename: str | None = None,
    ) -> RawDocument:
        path = Path(file_path).resolve()

        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        extension = path.suffix.lstrip(".").lower()

        if extension not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type '.{extension}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
            )

        document_name = original_filename if original_filename else path.name

        logger.info("Loading document: '%s' from path '%s' (type=%s)", document_name, path.name, extension)

        dispatch: dict[str, Any] = {
            "pdf": self._load_pdf,
            "docx": self._load_docx,
            "txt": self._load_txt,
            "xlsx": self._load_xlsx,
        }

        raw_doc: RawDocument = dispatch[extension](path, document_name)

        if not raw_doc.full_text.strip():
            raise ValueError(
                f"No text could be extracted from '{document_name}'. The file may be empty, image-only, or corrupted."
            )

        logger.info("Loaded '%s' — %d characters extracted.", document_name, len(raw_doc.full_text))
        return raw_doc

    def _load_pdf(self, path: Path, document_name: str) -> RawDocument:
        pages_text: list[str] = []
        page_char_offsets: list[int] = []
        cumulative = 0

        with pdfplumber.open(path) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages_text.append(text)
                page_char_offsets.append(cumulative)
                cumulative += len(text) + 1

        full_text = "\n".join(pages_text)

        metadata: dict[str, Any] = {
            "total_pages": total_pages,
            "page_char_offsets": page_char_offsets,
        }

        return RawDocument(
            document_name=document_name,
            file_type="pdf",
            full_text=full_text,
            metadata=metadata,
        )

    def _load_docx(self, path: Path, document_name: str) -> RawDocument:
        doc = DocxDocument(str(path))

        paragraphs: list[str] = [
            para.text for para in doc.paragraphs if para.text.strip()
        ]

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n".join(paragraphs)

        props = doc.core_properties
        metadata: dict[str, Any] = {
            "author": props.author or "",
            "created": props.created.isoformat() if props.created else "",
            "modified": props.modified.isoformat() if props.modified else "",
            "title": props.title or "",
            "paragraph_count": len(paragraphs),
        }

        return RawDocument(
            document_name=document_name,
            file_type="docx",
            full_text=full_text,
            metadata=metadata,
        )

    def _load_txt(self, path: Path, document_name: str) -> RawDocument:
        try:
            full_text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            logger.warning("'%s' is not valid UTF-8; falling back to Latin-1.", document_name)
            full_text = path.read_text(encoding="latin-1")

        metadata: dict[str, Any] = {
            "size_bytes": os.path.getsize(path),
            "line_count": full_text.count("\n"),
        }

        return RawDocument(
            document_name=document_name,
            file_type="txt",
            full_text=full_text,
            metadata=metadata,
        )

    def _load_xlsx(self, path: Path, document_name: str) -> RawDocument:
        wb = load_workbook(filename=str(path), read_only=True, data_only=True)
        sheet_texts: list[str] = []
        sheet_names: list[str] = wb.sheetnames

        for sheet_name in sheet_names:
            ws = wb[sheet_name]
            rows_text: list[str] = []

            for row in ws.iter_rows(values_only=True):
                non_empty = [str(cell) if cell is not None else "" for cell in row]
                if any(cell.strip() for cell in non_empty):
                    rows_text.append(" | ".join(non_empty))

            if rows_text:
                sheet_block = f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text)
                sheet_texts.append(sheet_block)

        wb.close()

        full_text = "\n\n".join(sheet_texts)

        metadata: dict[str, Any] = {
            "sheet_names": sheet_names,
            "sheet_count": len(sheet_names),
        }

        return RawDocument(
            document_name=document_name,
            file_type="xlsx",
            full_text=full_text,
            metadata=metadata,
        )
