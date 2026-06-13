"""
Unit tests for Phase 2 — Document Ingestion.

Tests cover:
  - DocumentLoader: all four file types + error cases.
  - DocumentChunker: chunk structure, metadata, edge cases.

Run with:
    pytest tests/test_ingestion.py -v

Fixtures create real temporary files on disk so tests exercise the
actual parsing logic rather than mocking internals.
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path

import pytest

from app.ingestion.chunker import DocumentChunker
from app.ingestion.loaders import DocumentLoader, UnsupportedFileTypeError
from app.models.responses import Chunk, RawDocument


# ======================================================================= #
# Helpers                                                                   #
# ======================================================================= #


def _make_txt_file(content: str, suffix: str = ".txt") -> Path:
    """Write content to a named temporary file and return its Path."""
    tmp = tempfile.NamedTemporaryFile(
        mode="w", suffix=suffix, delete=False, encoding="utf-8"
    )
    tmp.write(content)
    tmp.flush()
    tmp.close()
    return Path(tmp.name)


# ======================================================================= #
# DocumentLoader — TXT                                                      #
# ======================================================================= #


class TestTxtLoader:
    def test_loads_utf8_file(self) -> None:
        path = _make_txt_file("Hello, world!\nSecond line.")
        try:
            loader = DocumentLoader()
            doc = loader.load(path)
            assert isinstance(doc, RawDocument)
            assert doc.file_type == "txt"
            assert "Hello, world!" in doc.full_text
            assert doc.document_name == path.name
        finally:
            os.unlink(path)

    def test_metadata_contains_line_count(self) -> None:
        content = "Line 1\nLine 2\nLine 3\n"
        path = _make_txt_file(content)
        try:
            doc = DocumentLoader().load(path)
            assert "line_count" in doc.metadata
            assert doc.metadata["line_count"] == 3
        finally:
            os.unlink(path)

    def test_metadata_contains_size_bytes(self) -> None:
        content = "Some content"
        path = _make_txt_file(content)
        try:
            doc = DocumentLoader().load(path)
            assert "size_bytes" in doc.metadata
            assert doc.metadata["size_bytes"] > 0
        finally:
            os.unlink(path)

    def test_raises_on_empty_file(self) -> None:
        path = _make_txt_file("")
        try:
            with pytest.raises(ValueError, match="No text could be extracted"):
                DocumentLoader().load(path)
        finally:
            os.unlink(path)


# ======================================================================= #
# DocumentLoader — DOCX                                                     #
# ======================================================================= #


class TestDocxLoader:
    def test_loads_docx_file(self) -> None:
        from docx import Document as DocxDocument

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        try:
            doc = DocxDocument()
            doc.add_paragraph("First paragraph of the test document.")
            doc.add_paragraph("Second paragraph with more content.")
            doc.save(str(tmp_path))

            raw = DocumentLoader().load(tmp_path)
            assert raw.file_type == "docx"
            assert "First paragraph" in raw.full_text
            assert "Second paragraph" in raw.full_text
        finally:
            os.unlink(tmp_path)

    def test_docx_metadata_keys_present(self) -> None:
        from docx import Document as DocxDocument

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        try:
            doc = DocxDocument()
            doc.add_paragraph("Content for metadata test.")
            doc.save(str(tmp_path))

            raw = DocumentLoader().load(tmp_path)
            for key in ("author", "created", "modified", "title", "paragraph_count"):
                assert key in raw.metadata, f"Missing metadata key: {key}"
        finally:
            os.unlink(tmp_path)

    def test_docx_table_text_extracted(self) -> None:
        from docx import Document as DocxDocument

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        try:
            doc = DocxDocument()
            doc.add_paragraph("Intro paragraph.")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Header A"
            table.cell(0, 1).text = "Header B"
            table.cell(1, 0).text = "Value 1"
            table.cell(1, 1).text = "Value 2"
            doc.save(str(tmp_path))

            raw = DocumentLoader().load(tmp_path)
            assert "Header A" in raw.full_text
            assert "Value 2" in raw.full_text
        finally:
            os.unlink(tmp_path)


# ======================================================================= #
# DocumentLoader — XLSX                                                     #
# ======================================================================= #


class TestXlsxLoader:
    def test_loads_xlsx_file(self) -> None:
        from openpyxl import Workbook

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Sales"
            ws.append(["Product", "Revenue", "Units"])
            ws.append(["Widget A", 10000, 500])
            ws.append(["Widget B", 20000, 800])
            wb.save(str(tmp_path))

            raw = DocumentLoader().load(tmp_path)
            assert raw.file_type == "xlsx"
            assert "Product" in raw.full_text
            assert "Widget A" in raw.full_text
            assert "Sales" in raw.full_text
        finally:
            os.unlink(tmp_path)

    def test_xlsx_metadata_sheet_names(self) -> None:
        from openpyxl import Workbook

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        try:
            wb = Workbook()
            ws1 = wb.active
            ws1.title = "Sheet1"
            ws1.append(["Data"])
            ws2 = wb.create_sheet("Sheet2")
            ws2.append(["More Data"])
            wb.save(str(tmp_path))

            raw = DocumentLoader().load(tmp_path)
            assert "sheet_names" in raw.metadata
            assert "Sheet1" in raw.metadata["sheet_names"]
            assert "Sheet2" in raw.metadata["sheet_names"]
            assert raw.metadata["sheet_count"] == 2
        finally:
            os.unlink(tmp_path)


# ======================================================================= #
# DocumentLoader — Error cases                                              #
# ======================================================================= #


class TestLoaderErrors:
    def test_raises_on_missing_file(self) -> None:
        with pytest.raises(FileNotFoundError):
            DocumentLoader().load("/nonexistent/path/file.txt")

    def test_raises_on_unsupported_extension(self) -> None:
        path = _make_txt_file("content", suffix=".csv")
        try:
            with pytest.raises(UnsupportedFileTypeError):
                DocumentLoader().load(path)
        finally:
            os.unlink(path)


# ======================================================================= #
# DocumentChunker                                                           #
# ======================================================================= #


def _make_raw_doc(
    text: str,
    file_type: str = "txt",
    document_name: str = "test_doc.txt",
    metadata: dict | None = None,
) -> RawDocument:
    return RawDocument(
        document_name=document_name,
        file_type=file_type,
        full_text=text,
        metadata=metadata or {},
    )


class TestDocumentChunker:
    def test_returns_list_of_chunks(self) -> None:
        text = "Word " * 500
        raw = _make_raw_doc(text)
        chunks = DocumentChunker(chunk_size=500, chunk_overlap=50).chunk(raw)
        assert isinstance(chunks, list)
        assert all(isinstance(c, Chunk) for c in chunks)

    def test_chunk_count_is_reasonable(self) -> None:
        text = "The quick brown fox jumps over the lazy dog. " * 120
        raw = _make_raw_doc(text)
        chunks = DocumentChunker(chunk_size=1000, chunk_overlap=100).chunk(raw)
        assert 4 <= len(chunks) <= 10

    def test_chunk_fields_populated(self) -> None:
        text = "Sample content. " * 200
        raw = _make_raw_doc(text)
        chunks = DocumentChunker(chunk_size=500, chunk_overlap=50).chunk(raw)
        for chunk in chunks:
            assert chunk.chunk_id
            assert chunk.document_name == "test_doc.txt"
            assert chunk.chunk_text
            assert chunk.chunk_index >= 0

    def test_chunk_ids_are_unique(self) -> None:
        text = "Unique chunk test content. " * 300
        raw = _make_raw_doc(text)
        chunks = DocumentChunker(chunk_size=500, chunk_overlap=50).chunk(raw)
        ids = [c.chunk_id for c in chunks]
        assert len(ids) == len(set(ids)), "Duplicate chunk IDs detected"

    def test_chunk_indexes_are_sequential(self) -> None:
        text = "Sequential index test. " * 300
        raw = _make_raw_doc(text)
        chunks = DocumentChunker(chunk_size=500, chunk_overlap=50).chunk(raw)
        indexes = [c.chunk_index for c in chunks]
        assert indexes == list(range(len(chunks)))

    def test_metadata_inherited_from_raw_doc(self) -> None:
        text = "Metadata propagation test. " * 200
        raw = _make_raw_doc(text, metadata={"author": "Test Author", "line_count": 10})
        chunks = DocumentChunker(chunk_size=500, chunk_overlap=50).chunk(raw)
        for chunk in chunks:
            assert chunk.metadata.get("author") == "Test Author"
            assert chunk.metadata.get("line_count") == 10
            assert chunk.metadata.get("file_type") == "txt"

    def test_txt_chunks_have_no_page_number(self) -> None:
        text = "Text file content. " * 200
        raw = _make_raw_doc(text, file_type="txt")
        chunks = DocumentChunker(chunk_size=500, chunk_overlap=50).chunk(raw)
        assert all(c.page_number is None for c in chunks)

    def test_pdf_chunks_have_page_number(self) -> None:
        page1 = "Page one content. " * 30
        page2 = "Page two content. " * 30
        full_text = page1 + "\n" + page2
        page_char_offsets = [0, len(page1) + 1]

        raw = RawDocument(
            document_name="test.pdf",
            file_type="pdf",
            full_text=full_text,
            metadata={"total_pages": 2, "page_char_offsets": page_char_offsets},
        )
        chunks = DocumentChunker(chunk_size=300, chunk_overlap=30).chunk(raw)
        page_numbers = [c.page_number for c in chunks]
        assert all(p is not None for p in page_numbers)
        assert all(1 <= p <= 2 for p in page_numbers)

    def test_raises_on_empty_text(self) -> None:
        raw = _make_raw_doc("")
        with pytest.raises(ValueError, match="Cannot chunk"):
            DocumentChunker().chunk(raw)

    def test_chunk_batch_processes_multiple_docs(self) -> None:
        docs = [
            _make_raw_doc("Document one content. " * 100, document_name="doc1.txt"),
            _make_raw_doc("Document two content. " * 100, document_name="doc2.txt"),
        ]
        chunks = DocumentChunker(chunk_size=500, chunk_overlap=50).chunk_batch(docs)
        doc_names = {c.document_name for c in chunks}
        assert "doc1.txt" in doc_names
        assert "doc2.txt" in doc_names

    def test_chunk_batch_skips_empty_doc(self) -> None:
        docs = [
            _make_raw_doc("Valid content. " * 100, document_name="valid.txt"),
            _make_raw_doc("", document_name="empty.txt"),
        ]
        chunks = DocumentChunker(chunk_size=500, chunk_overlap=50).chunk_batch(docs)
        doc_names = {c.document_name for c in chunks}
        assert "valid.txt" in doc_names
        assert "empty.txt" not in doc_names

    def test_single_chunk_for_short_text(self) -> None:
        text = "Short document."
        raw = _make_raw_doc(text)
        chunks = DocumentChunker(chunk_size=1000, chunk_overlap=100).chunk(raw)
        assert len(chunks) == 1
        assert "Short document" in chunks[0].chunk_text